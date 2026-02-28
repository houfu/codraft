#!/usr/bin/env python3
"""Codraft document renderer.

Consolidates docx, html, and markdown rendering into a single CLI script.
Reads a Jinja2-style template and a JSON context file, renders the document,
optionally produces a PDF, validates the output, and prints a JSON result
to stdout.

Usage:
    python scripts/render.py \
        --template templates/_examples/invoice/invoice.html \
        --format html \
        --context context.json \
        --output-dir output/ \
        --job-name invoice_acme_2026-02-28 \
        --pdf
"""

import argparse
import json
import os
import re
import subprocess
import shutil
import sys


# ---------------------------------------------------------------------------
# Boolean coercion
# ---------------------------------------------------------------------------

_TRUE_STRINGS = {"true", "yes", "y"}
_FALSE_STRINGS = {"false", "no", "n"}


def coerce_booleans(context: dict) -> dict:
    """Convert string boolean representations to Python booleans.

    Operates recursively on nested dicts (e.g. inside loop items).
    """
    for key, value in context.items():
        if isinstance(value, str) and value.lower() in _TRUE_STRINGS:
            context[key] = True
        elif isinstance(value, str) and value.lower() in _FALSE_STRINGS:
            context[key] = False
        elif isinstance(value, list):
            for item in value:
                if isinstance(item, dict):
                    coerce_booleans(item)
    return context


# ---------------------------------------------------------------------------
# Job folder creation with deduplication
# ---------------------------------------------------------------------------


def create_job_dir(output_dir: str, job_name: str) -> tuple[str, str]:
    """Create a unique job directory under *output_dir*.

    If ``output_dir/job_name`` already exists, appends ``_2``, ``_3``, etc.

    Returns:
        A ``(job_dir, final_job_name)`` tuple.
    """
    base_job_name = job_name
    job_dir = os.path.join(output_dir, job_name)
    counter = 2
    while os.path.exists(job_dir):
        job_name = f"{base_job_name}_{counter}"
        job_dir = os.path.join(output_dir, job_name)
        counter += 1
    os.makedirs(job_dir)
    return job_dir, job_name


# ---------------------------------------------------------------------------
# Validation helpers
# ---------------------------------------------------------------------------

_PLACEHOLDER_RE = re.compile(r"\{\{\s*([\w.]+)\s*\}\}")
_CONTROLTAG_RE = re.compile(r"\{%.*?%\}")


def validate_text(text: str) -> dict:
    """Scan rendered text for unfilled placeholders and unprocessed tags."""
    unfilled = _PLACEHOLDER_RE.findall(text)
    tags = _CONTROLTAG_RE.findall(text)
    return {
        "passed": len(unfilled) == 0 and len(tags) == 0,
        "unfilled_variables": unfilled,
        "unprocessed_tags": tags,
    }


def validate_docx(docx_path: str) -> dict:
    """Open a rendered .docx with python-docx and scan for leftovers."""
    from docx import Document

    doc = Document(docx_path)
    all_text_parts: list[str] = []

    # Paragraphs
    for para in doc.paragraphs:
        all_text_parts.append(para.text)

    # Tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                all_text_parts.append(cell.text)

    # Headers and footers
    for section in doc.sections:
        for header_footer in (section.header, section.footer):
            if header_footer is not None:
                for para in header_footer.paragraphs:
                    all_text_parts.append(para.text)

    combined = "\n".join(all_text_parts)
    return validate_text(combined)


# ---------------------------------------------------------------------------
# Renderers
# ---------------------------------------------------------------------------


def render_docx(
    template_path: str,
    context: dict,
    job_dir: str,
    job_name: str,
    produce_pdf: bool,
) -> dict:
    """Render a .docx template using docxtpl."""
    from docxtpl import DocxTemplate

    doc = DocxTemplate(template_path)
    doc.render(context)

    docx_path = os.path.join(job_dir, f"{job_name}.docx")
    doc.save(docx_path)
    files = [docx_path]

    pdf_produced = False
    pdf_warning = None

    if produce_pdf:
        pdf_path = os.path.join(job_dir, f"{job_name}.pdf")

        # Attempt 1: docx2pdf
        try:
            import docx2pdf

            docx2pdf.convert(docx_path, pdf_path)
            pdf_produced = True
        except Exception as exc:
            pdf_warning = f"docx2pdf conversion failed: {exc}"

        # Attempt 2: LibreOffice headless
        if not pdf_produced:
            try:
                env = os.environ.copy()
                env["HOME"] = "/tmp"
                result = subprocess.run(
                    [
                        "soffice",
                        "--headless",
                        "--norestore",
                        "--nologo",
                        "--nofirststartwizard",
                        "--convert-to",
                        "pdf",
                        "--outdir",
                        job_dir,
                        docx_path,
                    ],
                    capture_output=True,
                    text=True,
                    timeout=60,
                    env=env,
                )
                if result.returncode == 0 and os.path.exists(pdf_path):
                    pdf_produced = True
                    pdf_warning = None
                else:
                    pdf_warning = (
                        f"LibreOffice failed (exit {result.returncode}): "
                        f"{result.stderr.strip()}"
                    )
            except FileNotFoundError:
                pdf_warning = (
                    "PDF conversion unavailable: neither docx2pdf nor "
                    "LibreOffice (soffice) found."
                )
            except subprocess.TimeoutExpired:
                pdf_warning = "LibreOffice conversion timed out after 60 seconds."
            except Exception as exc:
                pdf_warning = f"LibreOffice conversion failed: {exc}"

        if pdf_produced:
            files.append(pdf_path)

    validation = validate_docx(docx_path)

    return {
        "job_dir": job_dir,
        "files": files,
        "pdf_produced": pdf_produced,
        "pdf_warning": pdf_warning,
        "validation": validation,
    }


def render_html(
    template_path: str,
    context: dict,
    job_dir: str,
    job_name: str,
    produce_pdf: bool,
) -> dict:
    """Render an .html template using jinja2 + weasyprint."""
    from jinja2 import Template

    with open(template_path, encoding="utf-8") as f:
        template = Template(f.read())

    rendered_html = template.render(context)

    html_path = os.path.join(job_dir, f"{job_name}.html")
    with open(html_path, "w", encoding="utf-8") as f:
        f.write(rendered_html)
    files = [html_path]

    pdf_produced = False
    pdf_warning = None

    if produce_pdf:
        import weasyprint

        pdf_path = os.path.join(job_dir, f"{job_name}.pdf")
        weasyprint.HTML(filename=html_path).write_pdf(pdf_path)
        pdf_produced = True
        files.append(pdf_path)

    validation = validate_text(rendered_html)

    return {
        "job_dir": job_dir,
        "files": files,
        "pdf_produced": pdf_produced,
        "pdf_warning": pdf_warning,
        "validation": validation,
    }


def render_markdown(
    template_path: str,
    context: dict,
    job_dir: str,
    job_name: str,
    produce_pdf: bool,
) -> dict:
    """Render a .md template using jinja2, optionally convert to PDF."""
    import jinja2

    with open(template_path, "r", encoding="utf-8") as f:
        md_source = f.read()

    rendered_md = jinja2.Template(md_source).render(**context)

    md_path = os.path.join(job_dir, f"{job_name}.md")
    with open(md_path, "w", encoding="utf-8") as f:
        f.write(rendered_md)
    files = [md_path]

    pdf_produced = False
    pdf_warning = None

    if produce_pdf:
        pdf_path = os.path.join(job_dir, f"{job_name}.pdf")
        try:
            import markdown as md_lib
            import weasyprint

            html_content = md_lib.markdown(
                rendered_md, extensions=["tables", "fenced_code"]
            )
            weasyprint.HTML(string=html_content).write_pdf(pdf_path)
            pdf_produced = True
            files.append(pdf_path)
        except ImportError:
            pdf_warning = (
                "PDF not produced: install `markdown` and `weasyprint` "
                "to enable PDF output."
            )
        except Exception as exc:
            pdf_warning = f"PDF conversion failed: {exc}"

    validation = validate_text(rendered_md)

    return {
        "job_dir": job_dir,
        "files": files,
        "pdf_produced": pdf_produced,
        "pdf_warning": pdf_warning,
        "validation": validation,
    }


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

_RENDERERS = {
    "docx": render_docx,
    "html": render_html,
    "markdown": render_markdown,
}


def parse_args(argv: list[str] | None = None) -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Codraft document renderer — renders templates to docx/html/md with optional PDF."
    )
    parser.add_argument(
        "--template",
        required=True,
        help="Path to the template file (.docx, .html, or .md).",
    )
    parser.add_argument(
        "--format",
        required=True,
        choices=["docx", "html", "markdown"],
        help="Template format.",
    )
    parser.add_argument(
        "--context",
        required=True,
        help="Path to a JSON file containing the template context variables.",
    )
    parser.add_argument(
        "--output-dir",
        required=True,
        help="Base output directory (e.g. output/).",
    )
    parser.add_argument(
        "--job-name",
        required=True,
        help="Base name for the job folder and output files.",
    )
    parser.add_argument(
        "--pdf",
        action="store_true",
        default=False,
        help="Attempt PDF generation (soft-fail for docx and markdown).",
    )
    return parser.parse_args(argv)


def main(argv: list[str] | None = None) -> None:
    args = parse_args(argv)

    # Load context
    with open(args.context, "r", encoding="utf-8") as f:
        context = json.load(f)

    # Boolean coercion
    coerce_booleans(context)

    # Create deduplicated job directory
    job_dir, job_name = create_job_dir(args.output_dir, args.job_name)

    # Render
    renderer = _RENDERERS[args.format]
    result = renderer(
        template_path=args.template,
        context=context,
        job_dir=job_dir,
        job_name=job_name,
        produce_pdf=args.pdf,
    )

    # Ensure job_dir has a trailing slash for consistency
    result["job_dir"] = result["job_dir"].rstrip("/") + "/"

    # Output JSON result to stdout
    print(json.dumps(result, indent=2))


if __name__ == "__main__":
    main()
