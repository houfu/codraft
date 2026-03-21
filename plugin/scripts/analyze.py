"""CoQuill Analyzer — parse a template and generate manifest.yaml."""

import argparse
import os
import re
import sys

import yaml
from datetime import datetime, timezone


def detect_template(template_dir):
    """Find the template file and determine its format."""
    template_file = None
    template_format = None
    for f in os.listdir(template_dir):
        if f.endswith(".docx"):
            template_file = f
            template_format = "docx"
            break
        elif f.endswith(".html"):
            template_file = f
            template_format = "html"
            break
        elif f.endswith(".md"):
            template_file = f
            template_format = "markdown"
            break

    if not template_file:
        print("ERROR: No .docx, .html, or .md template found in", template_dir)
        sys.exit(1)

    return template_file, template_format


def check_cache(manifest_path, template_path, config_path):
    """Return True if the manifest needs regeneration."""
    if not os.path.exists(manifest_path):
        return True

    with open(manifest_path, "r") as mf:
        existing = yaml.safe_load(mf)

    if not existing or existing.get("schema_version", 0) < 2:
        return True

    analyzed_at = existing.get("analyzed_at", "")
    if not analyzed_at:
        return True

    analyzed_time = datetime.fromisoformat(analyzed_at.replace("Z", "+00:00"))
    tmpl_mtime = datetime.fromtimestamp(
        os.path.getmtime(template_path), tz=timezone.utc
    )
    config_mtime = None
    if os.path.exists(config_path):
        config_mtime = datetime.fromtimestamp(
            os.path.getmtime(config_path), tz=timezone.utc
        )

    if tmpl_mtime <= analyzed_time and (
        config_mtime is None or config_mtime <= analyzed_time
    ):
        return False

    return True


def extract_text(template_path, template_format):
    """Extract raw text from the template file."""
    if template_format == "docx":
        from docxtpl import DocxTemplate

        from docx import Document

        doc = DocxTemplate(template_path)
        try:
            text = doc.get_xml()
        except Exception:
            # Fallback: open with python-docx directly
            try:
                raw_doc = Document(template_path)
            except Exception:
                raw_doc = doc.docx
            if raw_doc is None:
                print(
                    f"WARNING: Could not read docx internals for {template_path}. "
                    "Falling back to python-docx Document()."
                )
                raw_doc = Document(template_path)
            text = ""
            for p in raw_doc.paragraphs:
                text += p.text + "\n"
            for table in raw_doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + "\n"
            for section in raw_doc.sections:
                if section.header:
                    for p in section.header.paragraphs:
                        text += p.text + "\n"
                if section.footer:
                    for p in section.footer.paragraphs:
                        text += p.text + "\n"
    else:
        with open(template_path, "r", encoding="utf-8") as f:
            text = f.read()
    return text


def analyze_template(text):
    """Two-pass analysis: find tags, classify variables into scopes."""
    # Regex patterns
    RE_IF = re.compile(r"\{%[-\s]*if\s+(.+?)\s*[-]?%\}")
    RE_ELSE = re.compile(r"\{%[-\s]*else\s*[-]?%\}")
    RE_ENDIF = re.compile(r"\{%[-\s]*endif\s*[-]?%\}")
    RE_FOR = re.compile(r"\{%[-\s]*for\s+(\w+)\s+in\s+(\w+)\s*[-]?%\}")
    RE_ENDFOR = re.compile(r"\{%[-\s]*endfor\s*[-]?%\}")
    RE_VAR = re.compile(r"\{\{\s*([\w.]+)\s*\}\}")
    RE_EQUALITY = re.compile(r"(\w+)\s*==\s*['\"](.+?)['\"]")

    # Find all tags with positions
    tags = []
    for m in RE_IF.finditer(text):
        tags.append(("if", m.start(), m.group(1).strip()))
    for m in RE_ELSE.finditer(text):
        tags.append(("else", m.start(), None))
    for m in RE_ENDIF.finditer(text):
        tags.append(("endif", m.start(), None))
    for m in RE_FOR.finditer(text):
        tags.append(("for", m.start(), (m.group(1), m.group(2))))
    for m in RE_ENDFOR.finditer(text):
        tags.append(("endfor", m.start(), None))
    for m in RE_VAR.finditer(text):
        tags.append(("var", m.start(), m.group(1)))

    tags.sort(key=lambda t: t[1])

    # Walk through tags and classify
    scope_stack = [{"type": "top"}]
    unconditional_vars = []
    conditionals = []
    loops = []
    condition_only_vars = set()

    for tag_type, pos, data in tags:
        current_scope = scope_stack[-1]

        if tag_type == "if":
            condition_expr = data
            eq_match = RE_EQUALITY.match(condition_expr)
            if eq_match:
                cond_block = {
                    "condition": condition_expr,
                    "gate_type": "equality",
                    "gate_variable": eq_match.group(1),
                    "gate_value": eq_match.group(2),
                    "if_variables": [],
                    "else_variables": [],
                }
            else:
                var_name = condition_expr.strip()
                cond_block = {
                    "condition": var_name,
                    "gate_type": "boolean",
                    "if_variables": [],
                    "else_variables": [],
                }
                condition_only_vars.add(var_name)

            conditionals.append(cond_block)
            scope_stack.append({"type": "if-branch", "block": cond_block})

        elif tag_type == "else":
            if len(scope_stack) > 1 and scope_stack[-1]["type"] == "if-branch":
                block = scope_stack[-1]["block"]
                scope_stack[-1] = {"type": "else-branch", "block": block}

        elif tag_type == "endif":
            if len(scope_stack) > 1 and scope_stack[-1]["type"] in (
                "if-branch",
                "else-branch",
            ):
                scope_stack.pop()

        elif tag_type == "for":
            loop_var, collection = data
            loop_block = {
                "loop_var": loop_var,
                "collection": collection,
                "variables": [],
            }
            loops.append(loop_block)
            scope_stack.append({"type": "loop-body", "block": loop_block})

        elif tag_type == "endfor":
            if len(scope_stack) > 1 and scope_stack[-1]["type"] == "loop-body":
                scope_stack.pop()

        elif tag_type == "var":
            var_name = data
            current = scope_stack[-1]

            if current["type"] == "top":
                if var_name not in unconditional_vars:
                    unconditional_vars.append(var_name)
            elif current["type"] == "if-branch":
                block = current["block"]
                if var_name not in [v["name"] for v in block["if_variables"]]:
                    block["if_variables"].append({"name": var_name})
            elif current["type"] == "else-branch":
                block = current["block"]
                if var_name not in [v["name"] for v in block["else_variables"]]:
                    block["else_variables"].append({"name": var_name})
            elif current["type"] == "loop-body":
                block = current["block"]
                loop_var = block["loop_var"]
                if "." in var_name and var_name.startswith(loop_var + "."):
                    sub_name = var_name[len(loop_var) + 1 :]
                    if sub_name not in [v["name"] for v in block["variables"]]:
                        block["variables"].append({"name": sub_name})
                else:
                    if var_name not in unconditional_vars:
                        unconditional_vars.append(var_name)

    # If a var appears at top-level AND in a conditional, keep it unconditional only
    for cond in conditionals:
        cond["if_variables"] = [
            v for v in cond["if_variables"] if v["name"] not in unconditional_vars
        ]
        cond["else_variables"] = [
            v for v in cond["else_variables"] if v["name"] not in unconditional_vars
        ]

    # Boolean inference: vars that appear only in {% if var %} and never in {{ var }}
    all_rendered_vars = set()
    for v in unconditional_vars:
        all_rendered_vars.add(v)
    for cond in conditionals:
        for v in cond["if_variables"]:
            all_rendered_vars.add(v["name"])
        for v in cond["else_variables"]:
            all_rendered_vars.add(v["name"])
    for loop in loops:
        for v in loop["variables"]:
            all_rendered_vars.add(v["name"])

    boolean_gate_vars = condition_only_vars - all_rendered_vars

    # Ensure gate variables are in unconditional list
    for cond in conditionals:
        if cond["gate_type"] == "boolean":
            gate_var = cond["condition"]
        else:
            gate_var = cond["gate_variable"]
        if gate_var not in unconditional_vars:
            unconditional_vars.append(gate_var)

    return unconditional_vars, conditionals, loops, boolean_gate_vars


def infer_type(name, boolean_gate_vars):
    if name in boolean_gate_vars:
        return "boolean"
    lower = name.lower()
    if lower.endswith(("_name", "_address")):
        return "text"
    if lower.endswith("_date"):
        return "date"
    if lower.endswith("_email"):
        return "email"
    if lower.endswith(("_amount", "_price", "_fee")):
        return "number"
    if lower.endswith(("_phone", "_tel", "_mobile")):
        return "phone"
    return "text"


def make_label(name):
    return name.replace("_", " ").title()


def build_var_entry(name, boolean_gate_vars, infer_name=None):
    return {
        "name": name,
        "label": make_label(name),
        "type": infer_type(infer_name or name, boolean_gate_vars),
    }


def build_manifest(
    template_dir, template_file, template_format, unconditional_vars,
    conditionals, loops, boolean_gate_vars, config_path,
):
    """Build the manifest dict from analyzed data and optional config."""
    # Build final variable lists
    final_variables = [build_var_entry(v, boolean_gate_vars) for v in unconditional_vars]

    for cond in conditionals:
        cond["if_variables"] = [
            build_var_entry(v["name"], boolean_gate_vars) for v in cond["if_variables"]
        ]
        cond["else_variables"] = [
            build_var_entry(v["name"], boolean_gate_vars) for v in cond["else_variables"]
        ]

    for loop in loops:
        loop["variables"] = [
            build_var_entry(
                v["name"], boolean_gate_vars, f"{loop['loop_var']}_{v['name']}"
            )
            for v in loop["variables"]
        ]
        loop["label"] = make_label(loop["collection"]).rstrip("s")
        loop["min_items"] = 1

    # Dependencies
    dependencies = {}
    for cond in conditionals:
        if cond["gate_type"] == "boolean":
            gate_var = cond["condition"]
        else:
            gate_var = cond["gate_variable"]

        dep_vars = [v["name"] for v in cond["if_variables"]] + [
            v["name"] for v in cond["else_variables"]
        ]
        if dep_vars:
            dependencies.setdefault(gate_var, []).extend(dep_vars)

    # Config merge
    config = None
    if os.path.exists(config_path):
        with open(config_path, "r", encoding="utf-8") as cf:
            config = yaml.safe_load(cf)

    if config and isinstance(config, dict):
        config_vars = config.get("variables", {})
        if config_vars:

            def merge_var_overrides(var_entry):
                """Merge config overrides into a variable entry."""
                name = var_entry["name"]
                if name in config_vars:
                    overrides = config_vars[name]
                    for field in [
                        "label",
                        "type",
                        "question",
                        "description",
                        "default",
                        "required",
                        "format_hint",
                        "choices",
                        "validation",
                    ]:
                        if field in overrides:
                            var_entry[field] = overrides[field]
                return var_entry

            final_variables = [merge_var_overrides(v) for v in final_variables]
            for cond in conditionals:
                cond["if_variables"] = [
                    merge_var_overrides(v) for v in cond["if_variables"]
                ]
                cond["else_variables"] = [
                    merge_var_overrides(v) for v in cond["else_variables"]
                ]
            for loop in loops:
                loop["variables"] = [
                    merge_var_overrides(v) for v in loop["variables"]
                ]

    # Count all unique variables
    all_var_names = set(v["name"] for v in final_variables)
    for cond in conditionals:
        for v in cond["if_variables"]:
            all_var_names.add(v["name"])
        for v in cond["else_variables"]:
            all_var_names.add(v["name"])
    for loop in loops:
        for v in loop["variables"]:
            all_var_names.add(v["name"])

    variable_count = len(all_var_names)

    # Build manifest
    manifest = {
        "schema_version": 2,
        "template": template_file,
        "template_path": os.path.join(template_dir, template_file),
        "format": template_format,
        "analyzed_at": datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ"),
        "variable_count": variable_count,
    }

    # Meta (from config only)
    if config and "meta" in config:
        manifest["meta"] = config["meta"]

    manifest["variables"] = final_variables
    manifest["conditionals"] = conditionals
    manifest["loops"] = loops
    manifest["dependencies"] = dependencies

    # Groups (from config only)
    if config and "groups" in config:
        manifest["groups"] = config["groups"]

    # Validation (from config only)
    if config and "validation" in config:
        manifest["validation"] = config["validation"]

    return manifest, variable_count


def main():
    parser = argparse.ArgumentParser(
        description="Analyze a CoQuill template and generate manifest.yaml"
    )
    parser.add_argument("template_dir", help="Path to the template directory")
    parser.add_argument(
        "--force",
        action="store_true",
        help="Force regeneration, skip cache check",
    )
    args = parser.parse_args()

    template_dir = args.template_dir

    # Step 1: Detect format
    template_file, template_format = detect_template(template_dir)
    template_path = os.path.join(template_dir, template_file)

    # Step 2: Check cache
    manifest_path = os.path.join(template_dir, "manifest.yaml")
    config_path = os.path.join(template_dir, "config.yaml")

    if not args.force:
        needs_regen = check_cache(manifest_path, template_path, config_path)
        if not needs_regen:
            print("Manifest is up to date. Skipping analysis.")
            sys.exit(0)

    # Step 3: Extract text
    text = extract_text(template_path, template_format)

    # Step 4: Analyze
    unconditional_vars, conditionals, loops, boolean_gate_vars = analyze_template(text)

    # Step 5–8: Build manifest
    manifest, variable_count = build_manifest(
        template_dir, template_file, template_format,
        unconditional_vars, conditionals, loops, boolean_gate_vars,
        config_path,
    )

    # Step 9: Write manifest
    with open(manifest_path, "w", encoding="utf-8") as mf:
        yaml.dump(
            manifest, mf, default_flow_style=False, sort_keys=False, allow_unicode=True
        )

    print(f"Manifest written: {manifest_path}")
    print(f"  Format: {template_format}")
    print(f"  Variables: {variable_count}")
    print(f"  Conditionals: {len(manifest['conditionals'])}")
    print(f"  Loops: {len(manifest['loops'])}")


if __name__ == "__main__":
    main()
